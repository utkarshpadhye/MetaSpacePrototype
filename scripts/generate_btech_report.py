from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
ASSET_DIR = OUT_DIR / "assets"
OUT_DOCX = OUT_DIR / "MetaSpace_BTech_Project_Report_2025_26.docx"

PROJECT_TITLE = "MetaSpace: A Spatial Virtual Office for Distributed Team Collaboration"
ACADEMIC_YEAR = "2025-2026"
DEPARTMENT = "Department of Computer Engineering and Technology"
UNIVERSITY = "MIT World Peace University, Kothrud, Pune 411 038, Maharashtra - India"


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name in ("Heading 1", "Heading 2", "Heading 3", "Title"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.bold = True
    styles["Title"].font.size = Pt(18)
    styles["Title"].font.bold = True


def para(doc: Document, text: str = "", align=None, bold: bool = False):
    p = doc.add_paragraph()
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    return p


def centered(doc: Document, text: str, size: int = 12, bold: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    return p


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.add_run(item).font.name = "Times New Roman"


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.add_run(item).font.name = "Times New Roman"


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, bold=True)
        shade_cell(hdr.cells[i], "D9EAF7")
        if widths:
            hdr.cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()
    return table


def draw_box(draw, xy, label, fill, outline="#1f2937"):
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    font = ImageFont.load_default(size=24)
    lines = label.split("\n")
    total = len(lines) * 30
    y = y1 + ((y2 - y1) - total) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2, y), line, fill="#111827", font=font)
        y += 30


def arrow(draw, start, end):
    draw.line([start, end], fill="#111827", width=4)
    ex, ey = end
    sx, sy = start
    dx = 1 if ex >= sx else -1
    draw.polygon([(ex, ey), (ex - 16 * dx, ey - 9), (ex - 16 * dx, ey + 9)], fill="#111827")


def make_diagrams() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture = ASSET_DIR / "metaspace_architecture.png"
    modules = ASSET_DIR / "metaspace_modules.png"

    img = Image.new("RGB", (1400, 850), "#ffffff")
    d = ImageDraw.Draw(img)
    draw_box(d, (80, 90, 390, 210), "Users\nBrowser Clients", "#dbeafe")
    draw_box(d, (540, 80, 890, 220), "React + TypeScript\nVite Frontend", "#dcfce7")
    draw_box(d, (1010, 80, 1320, 220), "Canvas World\nDOM Overlays", "#fef3c7")
    draw_box(d, (540, 350, 890, 500), "FastAPI Backend\nREST + WebSocket", "#fee2e2")
    draw_box(d, (80, 610, 390, 760), "PostgreSQL\nTenant Data", "#e0e7ff")
    draw_box(d, (540, 610, 890, 760), "Local AI STT\nCaptions + Summary", "#fce7f3")
    draw_box(d, (1010, 610, 1320, 760), "WebRTC\nAudio/Video", "#ccfbf1")
    arrow(d, (390, 150), (540, 150))
    arrow(d, (890, 150), (1010, 150))
    arrow(d, (715, 220), (715, 350))
    arrow(d, (540, 425), (390, 685))
    arrow(d, (715, 500), (715, 610))
    arrow(d, (890, 425), (1010, 685))
    d.text((80, 30), "Figure 5.1: Proposed MetaSpace high-level architecture", fill="#111827", font=ImageFont.load_default(size=28))
    img.save(architecture)

    img = Image.new("RGB", (1400, 850), "#ffffff")
    d = ImageDraw.Draw(img)
    draw_box(d, (80, 80, 360, 200), "World Renderer\nTiles, avatars", "#dbeafe")
    draw_box(d, (420, 80, 700, 200), "Communication\nChat, proximity", "#dcfce7")
    draw_box(d, (760, 80, 1040, 200), "Collaboration\nWhiteboard, docs", "#fef3c7")
    draw_box(d, (1080, 80, 1340, 200), "AI Assistant\nAna + STT", "#fce7f3")
    draw_box(d, (80, 360, 360, 500), "PM Module\nSprints, tasks,\nGantt, reports", "#fee2e2")
    draw_box(d, (420, 360, 700, 500), "CRM Module\nContacts,\ncompanies, deals", "#e0e7ff")
    draw_box(d, (760, 360, 1040, 500), "Access Control\nRoles, tenancy,\npermissions", "#ccfbf1")
    draw_box(d, (1080, 360, 1340, 500), "Notifications\nToasts, pulses,\nwebsocket", "#f3e8ff")
    draw_box(d, (420, 630, 1040, 760), "Shared Workspace Data Model\nWorkspace, users, projects, tasks, docs, CRM records", "#f8fafc")
    for x in (220, 560, 900, 1210):
        arrow(d, (x, 200), (720, 630))
    for x in (220, 560, 900, 1210):
        arrow(d, (x, 500), (720, 630))
    d.text((80, 30), "Figure 5.2: Functional modules and shared data model", fill="#111827", font=ImageFont.load_default(size=28))
    img.save(modules)
    return {"architecture": architecture, "modules": modules}


def title_page(doc: Document) -> None:
    centered(doc, "-", 12)
    para(doc)
    centered(doc, "Project Report", 16, True)
    centered(doc, "on", 12)
    centered(doc, PROJECT_TITLE, 18, True)
    para(doc)
    centered(doc, "Submitted by", 12, True)
    centered(doc, "Project Members", 12, True)
    centered(doc, "[Student Name 1] - [PRN Number]", 12)
    centered(doc, "[Student Name 2] - [PRN Number]", 12)
    centered(doc, "[Student Name 3] - [PRN Number]", 12)
    centered(doc, "[Student Name 4] - [PRN Number]", 12)
    para(doc)
    centered(doc, "Under the Internal Guidance of", 12, True)
    centered(doc, "Prof./Dr. [Project Guide Name]", 12)
    para(doc)
    centered(doc, "Under the External Guidance of (if applicable)", 12, True)
    centered(doc, "[External Guide Name / Not Applicable]", 12)
    para(doc)
    centered(doc, DEPARTMENT, 12, True)
    centered(doc, "MIT World Peace University, Kothrud,", 12)
    centered(doc, "Pune 411 038, Maharashtra - India", 12)
    centered(doc, ACADEMIC_YEAR, 12, True)
    doc.add_page_break()


def certificate(doc: Document) -> None:
    centered(doc, "DEPARTMENT OF COMPUTER ENGINEERING AND TECHNOLOGY", 14, True)
    para(doc)
    centered(doc, "C E R T I F I C A T E", 16, True)
    para(doc)
    para(
        doc,
        f'This is to certify that, [Student Names] of B.Tech. (Computer Science & Engineering) have completed their project titled "{PROJECT_TITLE}" and have submitted this Capstone Project Report towards fulfillment of the requirement for the Degree-Bachelor of Computer Science & Engineering (B.Tech-CSE) for the academic year {ACADEMIC_YEAR}.',
    )
    for _ in range(4):
        para(doc)
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r in table.rows:
        for cell in r.cells:
            cell.text = ""
    table.cell(0, 0).text = "[Prof./Dr. Guide Name]"
    table.cell(0, 1).text = "[Dr. Balaji M Patil]"
    table.cell(1, 0).text = "Project Guide"
    table.cell(1, 1).text = "Program Director"
    table.cell(2, 0).text = "DCET"
    table.cell(2, 1).text = "DCET"
    table.cell(3, 0).text = "MIT World Peace University, Pune"
    table.cell(3, 1).text = "MIT World Peace University, Pune"
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
    para(doc)
    para(doc, "Date:")
    doc.add_page_break()


def front_matter(doc: Document) -> None:
    centered(doc, "Acknowledgement", 16, True)
    para(
        doc,
        "We express our sincere gratitude to the Department of Computer Engineering and Technology, MIT World Peace University, for providing the opportunity and academic environment required to complete this capstone project. We are thankful to our project guide, Prof./Dr. [Project Guide Name], for continuous guidance, review, and encouragement throughout the development of MetaSpace.",
    )
    para(
        doc,
        "We also thank the faculty members, laboratory staff, and our peers for their feedback during requirement analysis, design validation, and prototype testing. Their suggestions helped us refine the project from a 2D virtual office prototype into a broader collaborative workspace with project management, CRM, documentation, voice, and AI-assisted meeting features.",
    )
    para(doc)
    para(doc, "Name of the Students")
    doc.add_page_break()

    centered(doc, "Abstract", 16, True)
    para(
        doc,
        "MetaSpace is a browser-based spatial virtual office prototype designed to make distributed collaboration more natural, contextual, and integrated. Conventional remote-work tools separate meetings, chat, task tracking, document collaboration, and customer management into different applications. This fragmentation increases context switching and reduces the sense of team presence. MetaSpace addresses this problem by combining a 2D top-down pixel-art workspace with proximity-based communication, interactive rooms, collaborative overlays, project management, CRM, documentation, notifications, and local AI-assisted meeting transcription.",
    )
    para(
        doc,
        "The project is implemented as a full-stack web application. The frontend uses React, TypeScript, Vite, Tailwind CSS, HTML5 Canvas, WebRTC hooks, and collaborative libraries such as Yjs and Tiptap. The backend uses FastAPI, SQLAlchemy, Alembic migrations, role-based access control, multi-tenant workspace isolation, REST APIs, and real-time notification support. The prototype includes avatar movement, room-based interactions, chat, proximity voice, conference-room workflows, screen sharing, whiteboard/document experiences, sprint boards, task lifecycle management, Gantt-style planning, CRM records, and local speech-to-text support.",
    )
    para(
        doc,
        "The outcome demonstrates that a spatial interface can serve not only as a communication layer but also as an operational workspace where users can discover teammates, enter context-aware conversations, and access project artifacts from the same virtual environment.",
    )
    doc.add_page_break()

    centered(doc, "List of Figures", 16, True)
    add_table(
        doc,
        ["Figure No.", "Figure Title", "Page No."],
        [["5.1", "Proposed MetaSpace high-level architecture", "18"], ["5.2", "Functional modules and shared data model", "19"]],
    )
    doc.add_page_break()
    centered(doc, "List of Tables", 16, True)
    add_table(
        doc,
        ["Table No.", "Table Title", "Page No."],
        [
            ["3.1", "Project scope, assumptions, and limitations", "9"],
            ["4.1", "Software and hardware requirements", "11"],
            ["4.2", "Risk management plan", "13"],
            ["7.1", "Representative test cases", "25"],
            ["10.1", "Project to outcome mapping", "31"],
        ],
    )
    doc.add_page_break()

    centered(doc, "Contents", 16, True)
    rows = [
        ["Abstract", "I"],
        ["List of Figures", "II"],
        ["List of Tables", "III"],
        ["Chapter 1 Introduction", "1"],
        ["Chapter 2 Literature Survey", "5"],
        ["Chapter 3 Problem Statement", "8"],
        ["Chapter 4 Project Requirements", "10"],
        ["Chapter 5 System Analysis and Proposed Architecture", "16"],
        ["Chapter 6 Project Plan", "22"],
        ["Chapter 7 Implementation", "24"],
        ["Chapter 8 Deployment and Security", "29"],
        ["Chapter 9 Applications", "30"],
        ["Chapter 10 Result and Analysis", "31"],
        ["Conclusion and Future Scope", "33"],
        ["References", "35"],
        ["Part B Individual Contribution", "36"],
        ["Appendices", "40"],
    ]
    add_table(doc, ["Section", "Page No."], rows)
    doc.add_page_break()


def part_a(doc: Document, diagrams: dict[str, Path]) -> None:
    centered(doc, "PART - A", 16, True)
    centered(doc, "GROUP / TEAM CONTRIBUTION", 14, True)
    doc.add_page_break()

    centered(doc, "Chapter 1", 14, True)
    centered(doc, "Introduction", 18, True)
    para(
        doc,
        "MetaSpace is a collaborative virtual workspace inspired by spatial office platforms. The project brings team presence, communication, and operational tools into a single browser-based environment. Users control avatars on a 2D tile map, move between rooms, interact with objects, and communicate through chat, proximity voice, conference calls, and screen sharing. The application is intended for distributed software teams, project groups, small organizations, and educational teams that need a lightweight but rich workspace for daily coordination.",
    )
    para(
        doc,
        "The need for this project arises from the limitations of conventional remote-work systems. Video meeting tools provide synchronous communication but do not create ambient awareness. Task management tools provide structure but do not show who is available for quick discussion. CRM and document tools store important work artifacts but remain disconnected from the communication space. MetaSpace integrates these workflows through a spatial metaphor: a project management room contains sprint boards and reports, a conference room supports richer calls and transcription, and workspace objects open the tools needed in that context.",
    )
    heading(doc, "1.1 Project Statement", 2)
    para(
        doc,
        "To design and implement a full-stack spatial virtual office prototype that enables distributed teams to collaborate through avatar-based presence, proximity communication, project management, CRM, documentation, notifications, and AI-assisted meeting support in a unified web application.",
    )
    heading(doc, "1.2 Area", 2)
    para(doc, "The project belongs to the areas of web engineering, collaborative systems, human-computer interaction, real-time communication, software project management, and enterprise productivity platforms.")
    heading(doc, "1.3 Aim", 2)
    bullets(
        doc,
        [
            "To provide a browser-based 2D virtual office where users can experience team presence through avatars and rooms.",
            "To reduce context switching by integrating chat, voice, meetings, tasks, documents, CRM, and notifications in one workspace.",
            "To implement secure multi-tenant backend APIs with role-based permissions and workspace isolation.",
            "To support project execution workflows such as sprint planning, task assignment, Gantt planning, reports, and personal work dashboards.",
        ],
    )
    heading(doc, "1.4 Implementation Overview", 2)
    para(
        doc,
        "The frontend is implemented using React and TypeScript with Vite as the build tool. HTML5 Canvas renders the world, avatars, tile layers, collision areas, and interaction hints. React DOM components render top bars, sidebars, modals, video overlays, project management panels, documentation overlays, CRM panels, notifications, and AI assistant panels. The backend is implemented using FastAPI with SQLAlchemy models and Alembic migrations. It exposes workspace, role, member, permission, project, sprint, milestone, task, documentation, and CRM routes. Real-time features are handled through browser channels, WebRTC hooks, and notification polling or websocket-style updates.",
    )
    heading(doc, "1.5 Applications of the Project", 2)
    bullets(
        doc,
        [
            "Remote and hybrid team collaboration.",
            "Academic project group coordination.",
            "Virtual office simulation for startups or small teams.",
            "Integrated project management and CRM workspace.",
            "Meeting transcription and summary support using local AI services.",
        ],
    )

    doc.add_page_break()
    centered(doc, "Chapter 2", 14, True)
    centered(doc, "Literature Survey", 18, True)
    para(
        doc,
        "The project is influenced by prior work in virtual offices, collaborative virtual environments, metaverse workplaces, WebRTC communication, and agile project management tools. Existing research and products show that remote collaboration succeeds when teams can maintain awareness, communicate quickly, and access shared artifacts without unnecessary friction.",
    )
    add_table(
        doc,
        ["Sr. No.", "Existing Work / System", "Key Idea", "Limitations Observed", "MetaSpace Approach"],
        [
            ["1", "Gather.town-style spatial offices", "2D avatar presence with proximity interaction", "Mostly communication-focused; external tools often needed for PM/CRM", "Adds PM, docs, CRM, roles, and AI meeting workflows inside the spatial office"],
            ["2", "Conventional video meeting tools", "Structured synchronous meetings", "No persistent spatial awareness or object-based context", "Uses conference rooms while retaining ambient avatar presence"],
            ["3", "Project management systems", "Boards, sprints, tasks, timelines", "Separated from communication layer", "Places sprint boards, backlog, timeline, and reports inside PM room overlays"],
            ["4", "Collaborative document platforms", "Shared editing and document storage", "Documents are not spatially discoverable", "Maps documents to workspace objects and overlay panels"],
            ["5", "CRM platforms", "Customer and deal tracking", "Usually separate from project execution and team presence", "Provides CRM access through role-restricted rooms and panels"],
        ],
    )
    heading(doc, "2.1 Earlier Limitations", 2)
    bullets(
        doc,
        [
            "Fragmented applications force users to move repeatedly between meetings, chat, task boards, documents, and CRM tools.",
            "Presence is usually binary, such as online/offline, and does not reflect spatial context or availability.",
            "Remote teams lack informal interactions that occur naturally in physical offices.",
            "Many collaboration prototypes focus on communication but do not address secure multi-tenant business workflows.",
        ],
    )
    heading(doc, "2.2 Proposed Improvement", 2)
    para(
        doc,
        "MetaSpace improves on these limitations by combining spatial presence with operational tools. The virtual map becomes a navigable interface to work. Users can enter rooms, approach teammates, interact with office objects, open project boards, manage tasks, review CRM data, join conference sessions, and receive notifications without leaving the environment.",
    )

    doc.add_page_break()
    centered(doc, "Chapter 3", 14, True)
    centered(doc, "Problem Statement", 18, True)
    para(
        doc,
        "Distributed teams require a platform that supports communication, project execution, shared documents, and customer workflows while preserving the social cues and informal discoverability of a physical office. Existing tools solve individual parts of this problem but do not provide an integrated, spatially meaningful workspace.",
    )
    heading(doc, "3.1 Scope, Assumptions and Limitations", 2)
    add_table(
        doc,
        ["Category", "Description"],
        [
            ["Scope", "2D virtual office, avatar movement, interactions, chat, proximity communication, conference support, PM module, docs module, CRM module, roles, permissions, notifications, and local STT service."],
            ["Assumptions", "Users access the system through modern browsers; backend APIs are deployed with a database; WebRTC permissions are granted by the browser; workspace data is scoped by workspace ID."],
            ["Limitations", "The current implementation is a prototype and uses local/demo identities in parts of the frontend. Production deployment would need full authentication, deployment hardening, and larger-scale load testing."],
        ],
    )
    heading(doc, "3.2 Project Objectives", 2)
    numbered(
        doc,
        [
            "Develop a React-based 2D virtual workspace with canvas rendering and interactive overlays.",
            "Implement secure backend APIs for workspace, member, role, permission, project, sprint, task, docs, and CRM workflows.",
            "Provide proximity-based communication and room-based collaboration workflows.",
            "Integrate project management functions including backlog, sprint board, task lifecycle, Gantt planning, reports, and notifications.",
            "Support local meeting transcription and summaries through a local AI service.",
        ],
    )

    doc.add_page_break()
    centered(doc, "Chapter 4", 14, True)
    centered(doc, "Project Requirements", 18, True)
    heading(doc, "4.1 Resources", 2)
    add_table(
        doc,
        ["Resource Type", "Details"],
        [
            ["Human Resources", "Frontend developer, backend developer, tester, documentation writer, project guide, and end users for validation."],
            ["Reusable Software Components", "React components, custom hooks, FastAPI routers, SQLAlchemy models, Alembic migrations, Yjs/Tiptap libraries, WebRTC hooks, local STT service, sprite and tile assets."],
            ["Hardware", "Development laptop/desktop with modern browser, microphone/camera for communication testing, local or hosted database server."],
            ["Software", "Node.js 20+, npm, Python 3.10+, FastAPI, SQLAlchemy, Alembic, PostgreSQL, Vite, React, TypeScript, Playwright, Pytest, Ruff, Tailwind CSS."],
        ],
    )
    heading(doc, "4.2 Requirements and Rationale", 2)
    add_table(
        doc,
        ["Requirement", "Rationale"],
        [
            ["Canvas-based world rendering", "Provides smooth 2D map, tile, avatar, and object rendering."],
            ["React overlay UI", "Allows rich panels, modals, sidebars, and controls over the canvas."],
            ["Workspace tenancy", "Prevents data leakage between organizations or project groups."],
            ["Role and permission system", "Supports Owner, Admin, Manager, Member, and Guest access levels."],
            ["Project management module", "Enables planning, task execution, sprints, milestones, and reports."],
            ["CRM module", "Allows contact, company, and deal tracking inside the workspace."],
            ["Local STT service", "Supports captions and meeting summaries without external transcription APIs."],
            ["Automated tests", "Validates security, routes, task lifecycle, notifications, and frontend behavior."],
        ],
    )
    heading(doc, "4.3 Risk Management", 2)
    add_table(
        doc,
        ["Risk Factor", "Impact", "Probability", "Mitigation"],
        [
            ["WebRTC browser permission failure", "Medium", "Medium", "Provide clear UI states and fallback chat/conference workflows."],
            ["Cross-workspace data leakage", "High", "Low", "Enforce workspace_id in backend dependencies, routes, and tests."],
            ["Performance drop during canvas rendering", "Medium", "Medium", "Use cached layers, limited redraw work, and performance baseline scripts."],
            ["Complex permission bugs", "High", "Medium", "Centralize check_permission and test precedence rules."],
            ["Incomplete production authentication", "High", "Medium", "Document prototype limitation and plan OAuth/JWT integration as future scope."],
        ],
    )
    heading(doc, "4.4 Functional Specifications", 2)
    bullets(
        doc,
        [
            "Graphical user interface with top bar, bottom bar, sidebars, overlays, modals, toast notifications, and avatar world.",
            "Internal interfaces between React components, hooks, canvas world state, and backend API adapters.",
            "External interfaces through REST endpoints, WebRTC media permissions, browser BroadcastChannel, and local STT HTTP service.",
            "Communication interfaces for chat, proximity voice, conference audio/video, screen sharing, and notifications.",
            "Security through role checks, permission overrides, guest restrictions, rate limits, and tenant-scoped queries.",
        ],
    )

    doc.add_page_break()
    centered(doc, "Chapter 5", 14, True)
    centered(doc, "System Analysis and Proposed Architecture", 18, True)
    para(
        doc,
        "The system follows a modular full-stack architecture. The frontend handles spatial interaction, user interface, media workflows, and local state. The backend provides persistent multi-tenant data, permission enforcement, REST APIs, and notification support. Optional local AI services provide speech-to-text functionality for conference-room captions and meeting summaries.",
    )
    doc.add_picture(str(diagrams["architecture"]), width=Inches(6.2))
    centered(doc, "Figure 5.1: Proposed MetaSpace high-level architecture", 11, True)
    para(doc)
    doc.add_picture(str(diagrams["modules"]), width=Inches(6.2))
    centered(doc, "Figure 5.2: Functional modules and shared data model", 11, True)
    heading(doc, "5.1 Design Considerations", 2)
    bullets(
        doc,
        [
            "Canvas is used for high-frequency world rendering, while DOM overlays are used for forms, panels, and collaboration tools.",
            "Backend mutation endpoints enforce permissions server-side because client-side checks are only for user experience.",
            "The data model uses workspace_id scoping to preserve tenant isolation.",
            "Rooms and map objects are used as spatial entry points for specialized tools.",
        ],
    )
    heading(doc, "5.2 Modules of the Project", 2)
    add_table(
        doc,
        ["Module", "Responsibility"],
        [
            ["World and Avatar Module", "Renders map, avatars, collisions, object interactions, proximity indicators, and room transitions."],
            ["Communication Module", "Supports chat messages, speech bubbles, proximity voice, conference calls, and screen sharing."],
            ["Project Management Module", "Manages projects, sprints, backlog, tasks, milestones, Gantt timeline, reports, and My Work dashboard."],
            ["Docs Module", "Provides project/workspace documentation and requirement knowledge base workflows."],
            ["CRM Module", "Manages contacts, companies, deals, and CRM reports."],
            ["Access Control Module", "Handles workspace roles, members, permission overrides, project role overrides, and guest restrictions."],
            ["AI Support Module", "Provides local STT captions and meeting summaries."],
        ],
    )
    heading(doc, "5.3 Low Level Design", 2)
    para(
        doc,
        "At low level, React components such as GameCanvas, TopBar, BottomBar, RightSidebar, ProjectManagementOverlay, DocsWorkspaceOverlay, CrmWorkspaceOverlay, ConferenceCallOverlay, ProximityVoiceOverlay, and AnaOverlay compose the interface. Hooks such as useGameLoop, useSessionPresence, useConferenceWebRTC, useProximityVoiceWebRTC, useScreenShare, useNotifications, useCalendarData, and useAnaAgent encapsulate repeated behavior. The backend uses route modules for workspaces, members, roles, overrides, projects, milestones, sprints, tasks, docs, and CRM. SQLAlchemy models define users, roles, workspace members, permission overrides, projects, sprints, milestones, tasks, notifications, docs, companies, contacts, and deals.",
    )

    doc.add_page_break()
    centered(doc, "Chapter 6", 14, True)
    centered(doc, "Project Plan", 18, True)
    add_table(
        doc,
        ["Phase", "Activities", "Deliverables"],
        [
            ["Phase 1", "Requirement analysis, study of spatial office workflows, frontend specification, project scope definition", "Requirement notes, frontend specification, report outline"],
            ["Phase 2", "Frontend canvas world, avatar movement, UI overlays, object interactions, chat", "Runnable Vite/React prototype"],
            ["Phase 3", "Backend bootstrap, migrations, workspace tenancy, role and permission system", "FastAPI backend, database models, tests"],
            ["Phase 4", "Project management, task lifecycle, notifications, reports, docs, CRM", "Operational PM/CRM/docs modules"],
            ["Phase 5", "Conference room, proximity voice, screen sharing, local STT, AI overlay", "Collaboration and meeting intelligence features"],
            ["Phase 6", "Testing, performance checks, report preparation, final review", "Validated prototype and project report"],
        ],
    )

    doc.add_page_break()
    centered(doc, "Chapter 7", 14, True)
    centered(doc, "Implementation", 18, True)
    heading(doc, "7.1 Methodology", 2)
    para(
        doc,
        "The project uses an incremental prototype methodology. The frontend world and interaction model were implemented first to establish the core user experience. Backend features were then added phase-by-phase, beginning with identity and tenancy, followed by project management, task lifecycle, documentation, CRM, notifications, and meeting-support modules. Each phase included validation through local builds, API tests, or end-to-end checks.",
    )
    heading(doc, "7.2 Algorithms and Logic", 2)
    bullets(
        doc,
        [
            "Avatar movement: keyboard input updates direction and target tile while collision maps prevent invalid movement.",
            "Proximity detection: distance between local avatar and peer avatars is computed each frame to activate nearby voice/video states.",
            "Permission resolution: revoke overrides are checked first, followed by grant overrides, project/workspace role permissions, and default deny.",
            "Sprint lifecycle: only one active sprint per project is allowed, incomplete tasks may be carried over, and reports use burndown snapshots.",
            "Dependency validation: task dependency cycles are rejected and timeline conflicts are highlighted.",
        ],
    )
    heading(doc, "7.3 Testing", 2)
    add_table(
        doc,
        ["Test Case No.", "Description", "Input", "Desired Output", "Result"],
        [
            ["TC-01", "Health endpoint", "GET /health", "Status 200 with app status", "Pass"],
            ["TC-02", "Permission precedence", "Role permission plus revoke override", "Access denied", "Pass"],
            ["TC-03", "Tenant isolation", "Requests using different workspace IDs", "No cross-workspace data", "Pass"],
            ["TC-04", "Sprint lifecycle", "Create sprint, add tasks, complete sprint", "Sprint completes with carry-over handling", "Pass"],
            ["TC-05", "Task notification", "Mention user in task comment", "Notification created and surfaced", "Pass"],
            ["TC-06", "Frontend build", "npm run build", "TypeScript and Vite build succeed", "Pass"],
            ["TC-07", "E2E app flow", "Open app and interact with workspace", "Core UI renders and interactions work", "Pass"],
        ],
    )
    heading(doc, "7.4 Performance Evaluation", 2)
    para(
        doc,
        "Performance considerations include canvas redraw efficiency, cached tile layers, controlled React state updates, rate limiting for backend routes, and benchmark scripts for frontend performance. The application separates high-frequency rendering from form-heavy overlays, which helps maintain responsiveness in the spatial interface.",
    )

    doc.add_page_break()
    centered(doc, "Chapter 8", 14, True)
    centered(doc, "Deployment Strategies and Security Aspects", 18, True)
    heading(doc, "8.1 Deployment Strategies", 2)
    para(
        doc,
        "The frontend can be built using npm run build and deployed as static assets through a web server or hosting platform. The FastAPI backend can be deployed behind an ASGI server such as Uvicorn/Gunicorn with PostgreSQL as the database. Alembic migrations should be executed during release. The local STT service may run as an optional internal service for environments where local captions are needed.",
    )
    heading(doc, "8.2 Security Aspects", 2)
    bullets(
        doc,
        [
            "All workspace data is scoped by workspace_id.",
            "Mutation routes enforce server-side permission checks.",
            "Guest users are prevented from receiving CRM permissions.",
            "Owner permissions are immutable and protected from deletion.",
            "Rate limiting protects the API from excessive global and mutation requests.",
            "Future production deployment should add full JWT/OAuth authentication, HTTPS, secure cookies, and audit logging.",
        ],
    )

    doc.add_page_break()
    centered(doc, "Chapter 9", 14, True)
    centered(doc, "Applications", 18, True)
    para(
        doc,
        "MetaSpace can be used by distributed teams that need informal communication, structured planning, and shared knowledge in one place. It can support software project teams, academic capstone groups, startup operations, virtual classrooms, remote support teams, and organizations that want a lightweight virtual-office experience with operational workflows.",
    )

    doc.add_page_break()
    centered(doc, "Chapter 10", 14, True)
    centered(doc, "Result and Analysis", 18, True)
    para(
        doc,
        "The implemented prototype demonstrates the feasibility of combining a spatial virtual office with project management, CRM, documentation, and meeting intelligence features. The frontend provides an interactive, game-like environment while the backend provides tenant-aware business logic. Testing notes in the repository indicate validation across permission rules, project APIs, phase 4 to phase 6 APIs, real-time notifications, security behavior, and frontend end-to-end flows.",
    )
    add_table(
        doc,
        ["Objective", "Outcome"],
        [
            ["Spatial virtual office", "Achieved through canvas map, avatar movement, rooms, objects, and overlays."],
            ["Communication features", "Implemented through chat, proximity voice hooks, conference WebRTC, screen sharing, and speech bubbles."],
            ["Operational tools", "Implemented through PM, docs, CRM, notifications, and task lifecycle modules."],
            ["Secure backend", "Implemented through FastAPI, role permissions, overrides, tenancy scoping, migrations, and tests."],
            ["AI meeting support", "Implemented through local STT server integration and Ana assistant overlay."],
        ],
    )
    heading(doc, "10.1 Project to Outcome Mapping", 2)
    add_table(
        doc,
        ["Sr. No.", "PRN No.", "Student Name", "Individual Project Student Specific Objective", "Learning Outcomes mapped"],
        [
            ["1", "[PRN]", "[Student Name]", "Frontend spatial office and interaction module", "Design and implement user-facing collaborative interfaces"],
            ["2", "[PRN]", "[Student Name]", "Backend tenancy, permissions, and API module", "Apply secure backend design and database modeling"],
            ["3", "[PRN]", "[Student Name]", "Project management, docs, CRM module", "Build integrated enterprise workflow modules"],
            ["4", "[PRN]", "[Student Name]", "Testing, deployment, and AI meeting support", "Validate full-stack quality and integration"],
        ],
    )

    doc.add_page_break()
    centered(doc, "Conclusion", 18, True)
    para(
        doc,
        "MetaSpace successfully presents a spatial collaboration prototype that moves beyond simple virtual meetings. It combines presence, proximity communication, room-based context, project execution, documentation, CRM, notifications, and local AI-assisted transcription into one web application. The project shows that a virtual office can be treated as an operating layer for teamwork rather than only as a visual meeting space.",
    )
    heading(doc, "Future Prospects", 2)
    bullets(
        doc,
        [
            "Production-grade authentication with OAuth/JWT and secure session handling.",
            "Scalable WebSocket and WebRTC signaling infrastructure.",
            "Persistent multiplayer synchronization across deployed workspaces.",
            "Advanced analytics for team productivity and meeting outcomes.",
            "More customizable maps, avatars, room templates, and organization branding.",
            "Deeper AI features such as action-item extraction, semantic document search, and meeting follow-up automation.",
        ],
    )

    doc.add_page_break()
    centered(doc, "References", 18, True)
    refs = [
        "M. Simova et al., Metaverse in the Virtual Workplace: Who and What is Driving the Remote Working Research, 2023.",
        "The Office of the Future: Virtual, Portable and Global.",
        "Virtual Office Prototype: A Unity-2D Work Simulation.",
        "Working Space: Virtual Office Prototype in Pandemic Era.",
        "YANA: Virtual Assistant to Support Home Office.",
        "MetaSpace project repository documentation: README.md, frontend-spec.md, crm-features.md, backend/README.md.",
        "FastAPI Documentation, SQLAlchemy Documentation, Alembic Documentation, React Documentation, Vite Documentation.",
        "WebRTC API documentation, Mozilla Developer Network.",
    ]
    for i, ref in enumerate(refs, 1):
        para(doc, f"[{i}] {ref}")


def part_b(doc: Document) -> None:
    doc.add_page_break()
    centered(doc, "PART - B", 16, True)
    centered(doc, "INDIVIDUAL CONTRIBUTION", 14, True)
    para(
        doc,
        "This section should be finalized after the exact student names, PRNs, and owned modules are confirmed. The following subsections are prepared in the required format and can be duplicated for each student.",
    )
    for idx, module in enumerate(
        [
            "Frontend Spatial Workspace Module",
            "Backend Access Control and API Module",
            "Project Management, Docs, and CRM Module",
            "Testing, Deployment, and AI Meeting Support Module",
        ],
        1,
    ):
        doc.add_page_break()
        centered(doc, f"Individual Contribution - Student {idx}", 16, True)
        para(doc, "Problem Statement: To contribute a defined module of the MetaSpace platform while aligning with the group project objectives and system architecture.")
        para(doc, "Name of the Student: [Student Name]")
        para(doc, "PRN No.: [PRN Number]")
        para(doc, f"Module Title: {module}")
        heading(doc, "Project's Module Objectives - Individual Perspective", 2)
        bullets(
            doc,
            [
                "Understand the module requirements and its relationship with the complete MetaSpace system.",
                "Design and implement the module using the technologies and patterns selected by the group.",
                "Validate the module through testing, integration, and documentation.",
            ],
        )
        heading(doc, "Project's Module Scope - Individual Perspective", 2)
        para(doc, "The scope includes module design, implementation, integration with related components, testing, and documentation. The module must remain consistent with the common architecture, data model, UI conventions, and security requirements.")
        heading(doc, "Project's Module(s) - Individual Contribution", 2)
        para(doc, f"The individual contribution covers {module}. Specific files, test cases, and implementation notes should be listed after final ownership confirmation.")
        heading(doc, "Hardware & Software Requirements", 2)
        para(doc, "Laptop/desktop system, modern browser, Node.js, npm, Python, FastAPI dependencies, PostgreSQL, Git, and development editor.")
        heading(doc, "Module Interfaces", 2)
        para(doc, "The module interfaces with React components, backend REST APIs, shared data models, hooks, route handlers, and/or database entities depending on its ownership.")
        heading(doc, "Module Dependencies", 2)
        para(doc, "Dependencies include shared workspace identity, role permissions, project data, UI state, backend API contracts, and validation tests.")
        heading(doc, "Module Design", 2)
        para(doc, "The module follows a componentized design. Responsibilities are separated between UI rendering, state management, API interaction, backend validation, persistence, and tests.")
        heading(doc, "Module Implementation", 2)
        para(doc, "Implementation details include development of the relevant frontend components or backend routes/models, connection to shared state/data, and integration with the main MetaSpace workflow.")
        heading(doc, "Module Testing Strategies", 2)
        para(doc, "Testing strategies include unit tests, integration tests, frontend end-to-end checks, manual UI verification, API route tests, permission tests, and regression validation.")
        heading(doc, "Module Deployment", 2)
        para(doc, "The module is deployed as part of the MetaSpace frontend build and/or FastAPI backend service. Database changes are deployed through Alembic migrations where applicable.")


def appendices(doc: Document) -> None:
    doc.add_page_break()
    centered(doc, "Appendices", 18, True)
    heading(doc, "A. Base Paper(s)", 2)
    para(doc, "Suggested base papers are listed in the References section. Final copies or links can be attached as required by the department.")
    heading(doc, "B. Plagiarism Report", 2)
    para(doc, "To be attached after running the final report through the approved plagiarism-checking tool.")
    heading(doc, "C. Formatting Notes", 2)
    add_table(
        doc,
        ["Item", "Applied / To be Checked"],
        [
            ["Main font", "Times New Roman 12"],
            ["Sub-heading font", "Times New Roman 14 Bold"],
            ["Line spacing", "1.5"],
            ["Text alignment", "Justified"],
            ["Title page and certificate", "No visible page number section in front matter; verify in Word before submission"],
            ["Missing personal details", "Student names, PRNs, guide name, signatures, and final page numbering"],
        ],
    )


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = make_diagrams()
    doc = Document()
    configure_document(doc)
    title_page(doc)
    certificate(doc)
    front_matter(doc)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.footer.is_linked_to_previous = False
    add_page_number(section.footer.paragraphs[0])
    part_a(doc, diagrams)
    part_b(doc)
    appendices(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
